from docx import Document

# Separate content into question-only and answer versions
content = """A1. In an (ε, δ)-DP system, a smaller ε usually means:
A. Stronger privacy, less noise
B. Stronger privacy, more noise
C. Weaker privacy, more noise
D. Unrelated to noise

A2. In the Laplace mechanism, the noise scale b is proportional to:
A. ε
B. Sensitivity (Δ)
C. Δ / ε
D. ε / Δ

A3. (True / False) 
If the same record is used in two independent ε = 0.5 queries, the total privacy loss is approximately ε = 1.0.

A4. Which operation directly reduces sensitivity Δ?
A. Decreasing ε
B. Adding random noise per record
C. Clipping each record’s contribution to a fixed range
D. Increasing sample size

A5. Compare two DP settings:  
P1 (ε = 1.0, δ = 1e-5) vs P2 (ε = 0.5, δ = 1e-6).  
Which provides stronger privacy?
A. P1
B. P2
C. Same
D. Cannot tell

A6. (Fill-in)
In the Gaussian mechanism, the standard deviation σ ≈ ____ × Δ / ε.

A7. The main difference between Local DP and Central DP is:
A. Noise added on server
B. Noise added on client
C. Both on server
D. Both on client

A8. Same count query, configuration X (ε = 1) vs Y (ε = 0.2). The expected absolute error of Y is:
A. Smaller
B. Larger
C. Same

A9. (True / False)  
If ε changes 1 → 0.5 and Δ changes 2 → 1, then b = Δ / ε stays the same (2 → 2).

A10. (Short answer) Why do we need privacy-budget management across multiple queries?
"""

answers = """Answers:

A1: B — Smaller ε → stronger privacy → more noise.
A2: C — b = Δ / ε.
A3: True — basic linear composition.
A4: C.
A5: B — smaller ε and δ.
A6: √(2 ln(1.25 / δ))
A7: B.
A8: B — smaller ε → more noise.
A9: True
A10: Because each query consumes part of the privacy budget; cumulative ε must stay limited to maintain overall privacy.
"""

# Create documents
doc_question = Document()
doc_question.add_heading("DP Quiz - Question Only Version", level=1)
doc_question.add_paragraph(content)
question_filename = "./test/DP_Quiz_Questions.docx"
doc_question.save(question_filename)

doc_answer = Document()
doc_answer.add_heading("DP Quiz - With Answers", level=1)
doc_answer.add_paragraph(content)
doc_answer.add_paragraph("\n---\n")
doc_answer.add_paragraph(answers)
answer_filename = "./test/DP_Quiz_Answers.docx"
doc_answer.save(answer_filename)

question_filename, answer_filename
